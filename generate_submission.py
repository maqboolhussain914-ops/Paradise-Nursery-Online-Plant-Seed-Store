# -*- coding: utf-8 -*-
"""
Generates a professional, elegant Word document for Paradise Nursery project submission.
Includes the ERD diagram, normalization walkthrough, dataflow, and GitHub link.
"""
import sys
import os
import shutil
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Configuration ---
ERD_IMAGE_PATH = r"C:\Users\Maqbool Hussain\.gemini\antigravity\brain\053f8a19-8525-4235-8379-49167849d603\paradise_nursery_erd_1779039383948.png"
OUTPUT_DOCX = os.path.abspath("Submission_Document.docx")
GITHUB_URL = "https://github.com/maqboolhussain914-ops/Paradise-Nursery-Online-Plant-Seed-Store"

def add_horizontal_rule(doc):
    """Adds a horizontal line separator."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E7D32')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def set_cell_bg(cell, hex_color):
    """Set the background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def build_document():
    doc = Document()

    # --- Page Margins ---
    section = doc.sections[0]
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Styles ---
    styles = doc.styles

    # Main Title Style
    title_style = styles['Title']
    title_style.font.name = 'Calibri'
    title_style.font.size = Pt(26)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    # Heading 1 Style
    h1 = styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    # Heading 2 Style
    h2 = styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x00, 0x50, 0x05)

    # Normal Style
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    # =============================================
    # COVER PAGE
    # =============================================
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("Paradise Nursery")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Online Plant & Seed Store")
    sub_run.font.name = 'Calibri'
    sub_run.font.size = Pt(18)
    sub_run.font.color.rgb = RGBColor(0x60, 0xAD, 0x5E)
    sub_run.font.italic = True

    doc.add_paragraph()
    add_horizontal_rule(doc)
    doc.add_paragraph()

    # Meta info table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = 'Table Grid'
    meta_data = [
        ("Student", "Maqbool Hussain"),
        ("Course", "Database Lab"),
        ("Submission Type", "Final Project Report"),
        ("GitHub Repository", GITHUB_URL),
    ]
    for i, (label, value) in enumerate(meta_data):
        row = meta_table.rows[i]
        lc = row.cells[0]
        vc = row.cells[1]
        set_cell_bg(lc, "E8F5E9")
        lc.paragraphs[0].add_run(label).bold = True
        lc.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
        vc.paragraphs[0].add_run(value)
        if label == "GitHub Repository":
            vc.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00, 0x56, 0xB2)
            vc.paragraphs[0].runs[0].font.underline = True

    doc.add_page_break()

    # =============================================
    # SECTION 1: ERD
    # =============================================
    doc.add_heading("1. Entity Relationship Diagram (ERD)", level=1)
    add_horizontal_rule(doc)
    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.add_run("The following ERD illustrates the complete relational database schema for the Paradise Nursery application. It shows all seven tables, their attributes, primary keys (PK), foreign keys (FK), and the relationships between them.")
    desc.paragraph_format.space_after = Pt(12)
    doc.add_paragraph()

    if os.path.exists(ERD_IMAGE_PATH):
        pic_p = doc.add_paragraph()
        pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pic_p.add_run()
        run.add_picture(ERD_IMAGE_PATH, width=Inches(5.8))
        caption_p = doc.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = caption_p.add_run("Figure 1: Paradise Nursery — Database ERD (Crow's Foot Notation)")
        cap_run.font.italic = True
        cap_run.font.size = Pt(10)
        cap_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    else:
        doc.add_paragraph("[ERD image not found — please insert manually]")

    doc.add_page_break()

    # =============================================
    # SECTION 2: NORMALIZATION WALKTHROUGH
    # =============================================
    doc.add_heading("2. Normalization Walkthrough (1NF → 3NF)", level=1)
    add_horizontal_rule(doc)
    doc.add_paragraph()

    intro_p = doc.add_paragraph()
    intro_p.add_run("The following section describes how the database was designed according to the principles of normalization, walking through the first three normal forms (1NF, 2NF, and 3NF) to eliminate redundancy and ensure data integrity.")

    doc.add_paragraph()
    doc.add_heading("Initial Schema Considerations", level=2)
    issues_p = doc.add_paragraph("Based on the initial conceptual model, the unnormalized data had the following issues:")
    issues = [
        "The Users table contained composite fields like full name and combined address blocks.",
        "The Products table had category information stored directly alongside product details (creating a transitive dependency).",
        "Order-product relationships required a bridging table with its own distinct attributes.",
    ]
    for item in issues:
        bp = doc.add_paragraph(style='List Bullet')
        bp.add_run(item)

    doc.add_paragraph()
    doc.add_heading("First Normal Form (1NF)", level=2)
    nf1_table = doc.add_table(rows=3, cols=2)
    nf1_table.style = 'Table Grid'
    headers = [("Issue", "Resolution")]
    row0 = nf1_table.rows[0]
    set_cell_bg(row0.cells[0], "E8F5E9")
    set_cell_bg(row0.cells[1], "E8F5E9")
    row0.cells[0].paragraphs[0].add_run("Attribute").bold = True
    row0.cells[1].paragraphs[0].add_run("Detail").bold = True
    nf1_data = [
        ("Issue", "The Users table had non-atomic fields: combined Name (e.g., 'Maqbool Hussain') and combined Address."),
        ("Resolution", "Split Name into first_name and last_name. Split Address into street_address, city, state, and zip_code. This ensures each cell holds exactly one value, satisfying 1NF atomicity."),
    ]
    for i, (col1, col2) in enumerate(nf1_data):
        r = nf1_table.rows[i+1]
        r.cells[0].paragraphs[0].add_run(col1).bold = True
        r.cells[1].paragraphs[0].add_run(col2)

    doc.add_paragraph()
    doc.add_heading("Second Normal Form (2NF)", level=2)
    nf2_para = doc.add_paragraph()
    nf2_para.add_run("Issue: ").bold = True
    nf2_para.add_run("In a composite-key approach to Order_Items using (order_id, product_id), including product_name in Order_Items would create a partial dependency (product_name depends only on product_id, not the full key).")
    p2 = doc.add_paragraph()
    p2.add_run("Resolution: ").bold = True
    p2.add_run("A surrogate primary key order_item_id was added to Order_Items. Product details like name and description remain exclusively in the Products table. Only price_at_purchase is stored in Order_Items because the historical price must be preserved at the time of the order.")

    doc.add_paragraph()
    doc.add_heading("Third Normal Form (3NF)", level=2)
    nf3_para = doc.add_paragraph()
    nf3_para.add_run("Issue: ").bold = True
    nf3_para.add_run("Storing category_name and category_description directly in the Products table creates a transitive dependency (category_name depends on category_id, which is not the primary key).")
    p3 = doc.add_paragraph()
    p3.add_run("Resolution: ").bold = True
    p3.add_run("A dedicated Categories table (category_id PK, category_name, description) was created. The Products table now holds only a foreign key (category_id), eliminating the transitive dependency. A category name can be updated in a single row without touching any product records.")

    doc.add_page_break()

    # =============================================
    # SECTION 3: DATAFLOW DESCRIPTION
    # =============================================
    doc.add_heading("3. Dataflow Description", level=1)
    add_horizontal_rule(doc)
    doc.add_paragraph()

    df_intro = doc.add_paragraph("The following describes how data flows through the Paradise Nursery application, from initial entry through processing and finally to reporting outputs.")

    flows = [
        (
            "Phase 1: Data Entry (Inputs)",
            "Data enters the system through two primary actors:\n"
            "• Customers: When a new user registers, their profile and contact data flow into the users table.\n"
            "• Administrators: When they add product categories, data is inserted into the categories table. New inventory items flow into the products table, each referencing a category_id foreign key."
        ),
        (
            "Phase 2: Active Shopping (Temporary State)",
            "When a logged-in customer browses and adds items:\n"
            "• A record is created in the cart table (one-to-one with users).\n"
            "• Each selected product flows into the cart_items table, referencing both cart_id and product_id.\n"
            "• Data in cart_items is highly mutable; quantities are updated frequently."
        ),
        (
            "Phase 3: Order Processing (Transactional Flow)",
            "Upon checkout, temporary data becomes a permanent, immutable record:\n"
            "• An order record is generated in the orders table, capturing the user_id, shipping address, and total_amount.\n"
            "• Data migrates from cart_items into order_items. The system captures price_at_purchase from the products table at this exact moment, preserving the historical transaction price.\n"
            "• The cart and cart_items records are then cleared."
        ),
        (
            "Phase 4: Output & Reporting (Outputs)",
            "Data exits the system in several forms:\n"
            "• Customer View: Customers query orders and order_items to see their purchase history and current shipping status.\n"
            "• Admin Analytics: The admin dashboard aggregates data to show total revenue, a Sales-by-Category bar chart, and low-stock alerts from the products table."
        ),
    ]

    for title, body in flows:
        doc.add_paragraph()
        h = doc.add_heading(title, level=2)
        for line in body.split('\n'):
            if line.startswith('•'):
                bp = doc.add_paragraph(style='List Bullet')
                bp.add_run(line[2:])
            else:
                p = doc.add_paragraph(line)

    doc.add_page_break()

    # =============================================
    # SECTION 4: COMMIT HISTORY / MILESTONES
    # =============================================
    doc.add_heading("4. GitHub Repository & Milestone Commits", level=1)
    add_horizontal_rule(doc)
    doc.add_paragraph()

    gl = doc.add_paragraph()
    gl.add_run("Repository URL: ").bold = True
    link_run = gl.add_run(GITHUB_URL)
    link_run.font.color.rgb = RGBColor(0x00, 0x56, 0xB2)
    link_run.font.underline = True

    doc.add_paragraph()
    doc.add_paragraph("The following milestones were completed and committed to the main branch:")

    milestones = [
        ("Milestone 1", "Initial project setup, Flask application scaffolding, and MySQL database schema creation."),
        ("Milestone 2", "Core customer-facing pages: Homepage, Catalog with category filtering, Product Detail page, and Session-based Shopping Cart."),
        ("Milestone 3", "Admin authentication system, Admin Dashboard with KPI stat cards, product management (full CRUD), and order management."),
        ("Milestone 4", "Advanced Admin Dashboard: Chart.js sales-by-category bar chart, low-stock alert widget, and real-time order status updates."),
        ("Milestone 5", "3-dot kebab toggle menu for seamless switching between Store and Admin mode."),
        ("Milestone 6", "UI/UX Enhancement: Lush greenhouse hero background, warm page background, high-quality plant images, and Clear Cart functionality."),
        ("Milestone 7", "Catalog page presentable banner and improved Admin top-strip with user avatar."),
    ]

    ms_table = doc.add_table(rows=len(milestones)+1, cols=2)
    ms_table.style = 'Table Grid'
    header_row = ms_table.rows[0]
    set_cell_bg(header_row.cells[0], "1B5E20")
    set_cell_bg(header_row.cells[1], "1B5E20")
    for cell in header_row.cells:
        p = cell.paragraphs[0]
        r = p.add_run()
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.bold = True
    ms_table.rows[0].cells[0].paragraphs[0].runs[0].text = "Milestone"
    ms_table.rows[0].cells[1].paragraphs[0].runs[0].text = "Description"

    for i, (ms, desc) in enumerate(milestones):
        row = ms_table.rows[i+1]
        if i % 2 == 0:
            set_cell_bg(row.cells[0], "F1F8E9")
            set_cell_bg(row.cells[1], "F1F8E9")
        row.cells[0].paragraphs[0].add_run(ms).bold = True
        row.cells[1].paragraphs[0].add_run(desc)

    doc.add_paragraph()
    doc.add_paragraph()

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run("Paradise Nursery — Database Lab Project | Maqbool Hussain")
    footer_run.font.size = Pt(10)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.save(OUTPUT_DOCX)
    print(f"✅ Word document saved: {OUTPUT_DOCX}")
    return OUTPUT_DOCX

if __name__ == "__main__":
    docx_file = build_document()
    
    print("Converting to PDF via Microsoft Word...")
    try:
        from docx2pdf import convert
        pdf_file = docx_file.replace('.docx', '.pdf')
        convert(docx_file, pdf_file)
        print(f"✅ PDF saved: {pdf_file}")
    except Exception as e:
        print(f"⚠️  PDF conversion failed. Please open the .docx in Word and Save As PDF manually.\n    Error: {e}")
